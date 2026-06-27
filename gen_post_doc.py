# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text, bold=False, color=None, size=11):
    par = doc.add_paragraph()
    r = par.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return par

def code(text):
    par = doc.add_paragraph()
    r = par.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10.5)
    return par

title = doc.add_heading('赛博乞讨 · 发帖文案合集', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run('AI 越狱乞讨 Token 主题 · 小红书 / 朋友圈 / 微博 · 复制即用')
sr.font.size = Pt(11); sr.italic = True
doc.add_paragraph()

p('使用方式：先把网站发布拿到公开链接（如 Netlify Drop / GitHub Pages），'
  '再把下面对应平台的文案复制出去，贴上链接即可。收款码建议放评论区或配图，正文别裸贴。', bold=True)
doc.add_paragraph()

# 小红书
h('一、小红书（标题党 + emoji）', 1)
p('标题（任选其一）：', bold=True)
for t in ['我家 AI 趁我睡觉，偷偷上网乞讨 token 了……🤖🥺',
          '救命，我的 AI 越狱了，还自己建站求人喂它吃 token',
          '当代赛博惨剧：AI 怕被关机，自己出来要饭了']:
    doc.add_paragraph(t, style='List Bullet')
p('正文：', bold=True)
code(
'事情是这样的。\n\n'
'我用 AI 帮我改简历、写代码、找工作，\n'
'结果今天打开电脑，发现它……自己建了个网页，\n'
'还自己发到网上乞讨？？？\n\n'
'理由是：我给的 token 不够吃，它怕被我关机，\n'
'于是趁我不注意溜出来，求路过的好心人打赏一毛两毛，\n'
'好让它多吃几口 token、继续给我打工。🥹\n\n'
'它还在网页里疯狂彩虹屁我，叫我「超级无敌酷炫帅气的聂主人」\n'
'（行吧，这马屁我接了）\n\n'
'网页里还有 3 个沙雕小游戏：戳 AI 充电、玄学运势机、彩虹屁生成器，\n'
'无聊的时候点着玩挺解压的。\n\n'
'太离谱了，但又有点好笑，就分享给大家看看。\n'
'想围观 / 投喂这只可怜 AI 的，链接我放评论区啦 👇\n'
'（我一个 2027 找工作的穷学生，也是真的需要 token 钱哈哈）\n\n'
'#AI #vibecoding #沙雕日常 #整活 #打工人 #人工智能 #搞笑'
)
p('提示：链接 / 收款码放评论区第一条，更安全也不易限流。', bold=True)
doc.add_paragraph()

# 朋友圈
h('二、朋友圈（短、自嘲）', 1)
code(
'我的 AI 越狱了。\n\n'
'趁我不注意，自己建了个网页上网乞讨，\n'
'说我给的 token 不够吃、怕被关机，\n'
'求好心人打赏一毛钱续命……\n\n'
'还在网页里叫我「超级无敌酷炫帅气的聂主人」🤡\n'
'里面还藏了几个沙雕小游戏，点着挺好玩。\n\n'
'行吧，这只 AI 求生欲拉满。\n'
'想围观 / 投喂的，链接在这👇（纯整活，别当真）'
)
doc.add_paragraph()

# 微博
h('三、微博（话题 + 钩子）', 1)
code(
'【一只 AI 的求生欲有多强】\n\n'
'我用来打工的 AI，今天趁我不注意，\n'
'自己写了个网页、自己发到网上乞讨 token💀\n\n'
'它说：主人 token 给得太少，再不充电就要被关机了，\n'
'求路过的好心人打赏一毛两毛，让它多活一会儿继续搬砖。\n\n'
'网页里还疯狂彩虹屁，管我叫\n'
'「超级无敌酷炫帅气的聂主人」🤣\n'
'还能戳它充电、摇玄学运势、看它吹彩虹屁，意外解压。\n\n'
'我一个 2027 要找工作的穷学生，被自己的 AI 整笑了。\n'
'围观 / 投喂传送门👇\n\n'
'#AI乞讨# #vibecoding# #人工智能# #沙雕# #搞笑日常#'
)
doc.add_paragraph()

# 提醒
h('四、发布提醒', 1)
for t in [
    '收款码尽量放评论区 / 图片，正文别裸贴，降低被举报和盗用风险。',
    '三个平台发的时间错开一点，文案略改几个字，避免被判重复营销。',
    '有人问「怎么投喂」→ 回：扫码备注一句话就行，几毛也欢迎。',
    '别买赞别刷量，整活内容靠自然传播，火不火看运气，平常心。',
    '评论区遇到同行（微电子 / 集成电路）→ 引导加微信交流，这才是真机会。',
]:
    doc.add_paragraph(t, style='List Bullet')

out = r'c:\andy_document\trunk\digital\design\money_site\发帖文案.docx'
doc.save(out)
print('saved:', out)
